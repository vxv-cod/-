'''https://docs-python.ru/packages/modul-python-docx-python/rabota-tablitsej/'''

'''Добавление и извлечение таблиц при помощи модуля python-docx.'''

'''Добавление табличных данных в документ DOCX.'''
# Пример добавления пустой таблицы, содержащей 2х2 ячейки:
from docx import Document
# создание пустого документа
doc = Document()
# добавляем пустую таблицу 2х2 ячейки
table = doc.add_table(rows=2, cols=2)

# Объект таблицы Table имеет несколько свойств и методов, которые необходимо вызвать, 
# чтобы заполнить таблицу данными. В качестве базового дальнейшего действия, всегда 
# можно получить доступ к ячейке таблицы, исходя из ее расположения в строке и столбце:
cell = table.cell(0, 1)

# Этот код возвратит объект ячейки Cell, которая расположена справа в верхней строке таблицы. 
# Обратите внимание, что индексы строк и столбцов начинаются с нуля, как в списке.
# В полученный объект ячейки можно записать какие-нибудь данные:

# добавляем данные как прогон абзаца 
# и выделяем текст жирным 
cell.paragraphs[0].add_run('Бык').bold = True
# можно записать данные в ячейку проще 
cell.text = 'Бык'
# что бы теперь отформатировать текст, нужно получить доступ к свойствам прогона ячейки
rc = cell.paragraphs[0].runs[0]
rc.font.name = 'Arial'
rc.font.bold = True

# Часто бывает проще получить доступ к ряду ячеек одновременно, например, 
# при заполнении таблицы переменной длины из источника данных. 
# Свойство таблицы Table.rows предоставляет доступ к отдельным строкам, 
# каждая из которых имеет свойство Table.rows[i].cells. Свойство .cells 
# как в строке, так и в столбце поддерживает доступ к ячейке по индексу (как с списку):
# получаем 2-ю строку таблицы
row = table.rows[1]
# запись данных в ячейки
row.cells[0].text = 'Заяц'
row.cells[1].text = 'Волк'

# Последовательности Table.rows и Table.columns в таблице являются итерируемыми, 
# следовательно можно использовать их непосредственно в цикле for. То же самое 
# с последовательностями ячеек, например для первой строки таблицы 
# Table.rows[0].cells или для первого столбца Table.columns[0].cells:
# читаем ячейки таблицы `table`
for row in table.rows:
    for cell in row.cells:
        print(cell.text)

# Если необходимо узнать количество строк или столбцов в таблице, 
# то просто используйте функцию len() для соответствующей последовательности:
# количество строк в таблице
row_count = len(table.rows)
# количество колонок в таблице
col_count = len(table.columns)

# Также можно добавлять строки в таблицу постепенно, например:
row = table.add_row()


# Это может быть очень удобно для построения таблицы переменной длины:
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
# создание пустого документа
doc = Document()
# данные таблицы без названий колонок
items = (
    (7, '1024', 'Плюшевые котята'),
    (3, '2042', 'Меховые пчелы'),
    (1, '1288', 'Ошейники для пуделей'),
)
# добавляем таблицу с одной строкой 
# для заполнения названий колонок
table = doc.add_table(1, len(items[0]))
# определяем стиль таблицы
table.style = 'Light Shading Accent 1'
# Получаем строку с колонками из добавленной таблицы 
head_cells = table.rows[0].cells
# добавляем названия колонок (enumerate - перечислять)
for i, item in enumerate(['Кол-во', 'ID', 'Описание']):
    p = head_cells[i].paragraphs[0]
    # название колонки
    p.add_run(item).bold = True
    # выравниваем посередине
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# добавляем данные к существующей таблице
for row in items:
    # добавляем строку с ячейками к объекту таблицы
    cells = table.add_row().cells
    for i, item in enumerate(row):
        # вставляем данные в ячейки
        cells[i].text = str(item)
        # если последняя ячейка
        if i == 2:
            # изменим шрифт
            cells[i].paragraphs[0].runs[0].font.name = 'Arial'
doc.save('test.docx')

# То же самое работает для столбцов, хотя строить таблицу таким способом не удобно.
# MS Word имеет набор предварительно отформатированных стилей таблиц, которые можно 
# выбрать из его галереи стилей таблиц. Применить один из них к таблице можно следующим образом:
table.style = 'Light Shading Accent 1'

# Обратите внимание, что имя стиля таблицы немного отличается от имени, 
# отображаемого в пользовательском интерфейсе MS Word. Дефис, если он есть, 
# то его необходимо удалить. Например, Light Shading - Accent 1 становится Light Shading Accent 1.

# Чтобы узнать название стиля таблицы, наведите указатель мыши на его эскиз в галерее стилей таблиц Word.

# Важно!!! Встроенные стили хранятся в файле WordprocessingML под своим английским именем, 
# например 'Table Grid', и не зависят от локализации MS Word. Так как модуль python-docx 
# работает с файлом WordprocessingML, то поиск стиля должен использовать английское имя. 
# Если файл WordprocessingML не найден (MS Word не установлен, например в OS Linux) то 
# модуль python-docx работает со своей версией этого файла. Что бы создать сопоставление 
# между именами стилей на русском языке и именами на английском языке посетите эту ссылку.

'''Все стили таблиц можно посмотреть, выполнив код:'''
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
doc = Document()
all_styles = doc.styles
table_styles = [s for s in all_styles if s.type == WD_STYLE_TYPE.TABLE]
for style in table_styles:
    print(table_styles.name)
# Normal Table
# Table Grid
# Light Shading
# Light Shading Accent 1


'''Извлечение табличных данных их документов DOCX.'''
# При чтении существующего документа DOCX, все находящиеся в нем объекты таблиц Table 
# группируются в последовательности Document.tables. Следовательно, что бы узнать 
# количество таблиц в документе, нужно вызвать функцию len() для этой последовательности.
# Доступ к объектам таблиц будем осуществлять по индексу последовательности Document.tables.
# Смотрим пример:

from docx import Document

doc = Document('test.docx')
# последовательность всех таблиц документа
all_tables = doc.tables
print('Всего таблиц в документе:', len(all_tables))
# создаем пустой словарь под данные таблиц
data_tables = {i:None for i in range(len(all_tables))}
# проходимся по таблицам
for i, table in enumerate(all_tables):
    print('\nДанные таблицы №', i)
    # создаем список строк для таблицы `i` (пока пустые)
    data_tables[i] = [[] for _ in range(len(table.rows))]
    # проходимся по строкам таблицы `i`
    for j, row in enumerate(table.rows):
        # проходимся по ячейкам таблицы `i` и строки `j`
        for cell in row.cells:
            # добавляем значение ячейки в соответствующий
            # список, созданного словаря под данные таблиц
            data_tables[i][j].append(cell.text)

    # смотрим извлеченные данные 
    # (по строкам) для таблицы `i`
    print(data_tables[i])
    print('\n')

print('Данные всех таблиц документа:')
print(data_tables)


'''https://docs-python.ru/packages/modul-python-docx-python/obekt-table/'''

'''Свойства и методы объекта Table.'''
table = doc.add_table(rows=2, cols=2)

# добавляет колонку
table.add_column()
# добавляет строку
table.add_row()
# выравнивание таблицы
table.alignment
# авто подгонка ширины колонок таблицы
table.autofit
# возвращает экземпляр определенной ячейки
table.cell() 
# список ячеек в определенном столбце
table.column_cells()
# список столбцов таблицы
table.columns
# список ячеек в определенной строке
table.row_cells()
# список строк таблицы
table.rows
# стиль таблицы
table.style
# направление, в котором упорядочены ячейки таблицы
table.table_direction

'''Table.add_column(width):'''
# Метод Table.add_column() добавляет колонку шириной width в крайний правый угол таблицы и возвращает экземпляр Column добавленной колонки.
'''# Объект Column имеет свойства:'''
# cells - последовательность экземпляров Cell, соответствующих ячейкам в этом столбце.
# table - ссылка на объект таблицы Table, которому принадлежит этот столбец.
# width - ширина этого столбца в величине EMU или None, если ширина явно не задана.

'''# Table.add_row():'''
# Метод Table.add_row() добавляет строку в самую нижнюю часть таблицы и возвращает экземпляр Row, только что добавленной строки.
'''Объект {Row} имеет свойства:'''
# cells - последовательность экземпляров Cell, соответствующих ячейкам в этом столбце.
# table - ссылка на объект таблицы Table, которому принадлежит этот столбец.
# height - возвращает/устанавливает объект Length, представляющий высоту этой ячейки, или None, если высота явно не задана.
# height_rule - возвращает/устанавливает правило высоты этой ячейки как член перечисления WD_ROW_HEIGHT_RULE или None, если явное правило высоты не установлено.

'''Table.alignment:'''
# Свойство Table.alignment возвращает/устанавливает правило, которое определяет расположение этой таблицы 
# между полями страницы. Значение является членом перечисления WD_TABLE_ALIGNMENT или None. 
# Значение None говорит о том, что действующее значение наследуется от иерархии стилей.

'''Table.autofit:'''
# Свойство Table.autofit принимает/возвращает True, если ширина столбцов автоматически регулируется для лучшего соответствия содержимому ячеек и False, если макет таблицы фиксированный.
# Ширина столбца корректируется в любом случае, если общая ширина столбцов превышает ширину страницы.

'''Table.cell(row_idx, col_idx):'''
# Метод Table.cell() возвращает экземпляр Сell, соответствующий ячейке таблицы на пересечении row_idx, col_idx, где (0, 0) является верхней, самой левой ячейкой.

'''Table.column_cells(column_idx):'''
# Метод Table.column_cells() представляет собой последовательность ячеек Сell в столбце с номером column_idx в этой таблице.

'''Table.columns:'''
# Свойство Table.columns представляет собой последовательность объектов столбцов Column в этой таблице.

'''Table.row_cells(row_idx):'''
# Метод Table.row_cells() представляет собой последовательность ячеек Сell в строке с 
# номером row_idx в этой таблице.
# Пустые ячейки не заполняют список

'''Table.rows:'''
# Свойство Table.rows представляет собой последовательность строк Row в этой таблице. 
# Поддерживает функцию len(), итерацию, доступ к строке по индексу, а так же получение среза строк.

'''Table.style:'''
# Свойство Table.style устанавливает/возвращает объект стиля таблицы Style или Имя Стиля Таблицы, которое встроено в пользовательский интерфейс редактора MS Word.
# Если таблица не имеет стиля, то возвращается стиль таблицы по умолчанию для документа (часто обычная таблица). Назначение этому свойству значения None удаляет любой применяемый напрямую стиль таблицы, заставляя его наследовать стиль таблицы документа по умолчанию.
# Обратите внимание, что имя стиля таблицы немного отличается от имени, отображаемого в пользовательском интерфейсе MS Word. Дефис, если он есть, то его необходимо удалить. Например, Light Shading - Accent 1 становится Light Shading Accent 1.
# Важно!!! Встроенные стили хранятся в файле WordprocessingML под своим английским именем, например 'Table Grid', и не зависят от локализации MS Word. Так как модуль python-docx работает с файлом WordprocessingML, то поиск стиля должен использовать английское имя. 
# Если файл WordprocessingML не найден (MS Word не установлен, например в OS Linux) то модуль python-docx работает со своей версией этого файла. Что бы создать сопоставление между именами стилей на русском языке и именами на английском языке посетите эту ссылку.

'''Table.table_direction:'''
# Свойство Table.table_direction это элемент WD_TABLE_DIRECTION, который указывает направление, 
# в котором упорядочены ячейки таблицы, например, WD_TABLE_DIRECTION.LTR или WD_TABLE_DIRECTION.RTL.
# Значение None указывает, что значение наследуется от иерархии стилей.
from docx.enum.table import WD_TABLE_DIRECTION
# создание документа
document = Document()
table = document.add_table(3, 3)
table.direction = WD_TABLE_DIRECTION.RTL


'''Свойства и методы объекта ячейки таблицы Cell.'''
Cell = table.cell(0, 1)
# добавляет абзац в ячейку,
Cell.add_paragraph() 
# добавляет таблицу в ячейку,
Cell.add_table()
# объединяет ячейки в одну,
Cell.merge()
# список объектов абзацев в ячейке,
Cell.paragraphs
# список таблиц в ячейке,
Cell.tables
# ВСЕ содержимое ячейки в виде строки текста,
Cell.text
# вертикальное выравнивание,
Cell.vertical_alignment
# ширина ячейки,
Cell.width


Cell.add_paragraph(text='', style=None)
# Метод Cell.add_paragraph() возвращает недавно добавленный абзац Paragraph в конец содержимого этой ячейки. 
# Если присутствует текст text, то он добавляется к абзацу за один прогон Run. Если указан style, 
# то применяется стиль для этого абзаца. Если стиль не указан или имеет значение None, то результат 
# будет таким, как если бы был применен стиль с именем 'Normal'.
# Обратите внимание, что на форматирование текста в ячейке может влиять стиль таблицы. 
# Аргумент text может содержать символы табуляции \t, которые преобразуются в соответствующую XML-форму. 
# Текст также может включать символы новой строки \n или возврата каретки \r, каждый из которых преобразуется в разрыв строки.

Cell.add_table(rows, cols)
# Метод Cell.add_table() возвращает недавно добавленную таблицу Paragraph в конец содержимого этой ячейки. 
# После таблицы добавляется пустой абзац, так как спецификация MS Word требует, чтобы последним элементом в каждой ячейке был абзац.

Cell.merge(other_cell)
# Метод Cell.merge() возвращает объединенную ячейку, созданную путем охвата прямоугольной области, 
# в которой эта ячейка и другая ячейка other_cell являются диагональными углами. Вызывает ошибку InvalidSpanError, 
# если ячейки не определяют прямоугольную область.

Cell.paragraphs
# Свойство Cell.paragraphs представляет собой список объектов абзацев в ячейке. 
# Ячейка таблицы должна содержать по крайней мере один элемент уровня блока и заканчиваться абзацем. 
# По умолчанию новая ячейка содержит один абзац. Свойство только для чтения.

Cell.tables
# Свойство Cell.tables представляет собой список объектов таблиц в ячейке в порядке их появления. 
# Свойство только для чтения.

Cell.text
# Свойство Cell.text представляет собой ВСЕ содержимое этой ячейки в виде строки текста. 
# Назначение строки этому свойству заменяет все существующее содержимое одним абзацем Paragraph, 
# содержащим назначенный текст, за один прогон Run.

Cell.vertical_alignment
# Свойство Cell.vertical_alignment возвращает/устанавливает вертикальное выравнивание 
# ячейки как член перечисления WD_CELL_VERTICAL_ALIGNMENT или None.
# Значение None указывает, что вертикальное выравнивание для этой ячейки унаследовано. 
# Назначение None приводит к удалению явно определенного вертикального выравнивания и восстановлению наследования.

Cell.width
# Свойство Cell.width возвращает/устанавливает ширину этой ячейки в величине EMU или None, если ширина явно не задана.



